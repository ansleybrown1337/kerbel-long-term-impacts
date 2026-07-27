from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
OUT = REPO / "docs" / "drafts" / "Chapter4_Results_Discussion_Limitations_Draft_v2.docx"
RESULTS = REPO / "results" / "comparison" / "v3p0_physical_event"
FIGURES = REPO / "figures" / "comparison" / "v3p0_physical_event"

ANALYTE_ORDER = ["NH4", "NO2", "NO3", "OP", "Se", "TDS", "TKN", "TN", "TP", "TSS"]
TREATMENT_ORDER = ["CT", "MT", "ST"]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=55, bottom=55, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        keep_row_together(row)
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_section_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)


def set_section_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def configure_styles(doc):
    set_section_portrait(doc.sections[0])
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for name, size, before, after in (
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 10, 4),
        ("Heading 3", 12, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_together = True


def add_body(doc, text, highlight=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    run = paragraph.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraph


def add_mixed_body(doc, parts):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    for text, highlight in parts:
        run = paragraph.add_run(text)
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return paragraph


def add_review_note(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(
        "Draft review note. Yellow highlighting identifies literature connections or "
        "substantive interpretation that should receive explicit author review."
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_title(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("Chapter 4 Results, Discussion, and Limitations")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run("Bayesian and ML reconstruction of the long-term tillage dataset")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_figure(doc, number, path, caption_text, alt_text, width):
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set("name", f"Figure {number}")
    inline.docPr.set("descr", alt_text)

    caption = doc.add_paragraph(style="Caption")
    label = caption.add_run(f"Figure {number}. ")
    label.bold = True
    caption.add_run(caption_text)

    source = doc.add_paragraph()
    source.paragraph_format.line_spacing = 1.0
    source.paragraph_format.space_after = Pt(6)
    source.paragraph_format.keep_together = True
    run = source.add_run(f"Draft source: {path.relative_to(REPO).as_posix()}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_table_caption(doc, number, title):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.paragraph_format.keep_with_next = True
    label = paragraph.add_run(f"Table {number}. ")
    label.bold = True
    paragraph.add_run(title)


def set_table_font(table, size):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(size)
                    if row_index == 0:
                        run.bold = True


def add_table(doc, headers, rows, widths, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], "D9E2F3")
    set_repeat_table_header(table.rows[0])

    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, "F7F7F7")

    set_table_widths(table, widths)
    set_table_font(table, font_size)
    return table


def add_table_note(doc, text, source_path=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(f"Note. {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    if source_path is not None:
        run = paragraph.add_run(f" Draft source: {source_path.relative_to(REPO).as_posix()}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_landscape(section)
    return section


def add_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_portrait(section)
    return section


def pct(value):
    return f"{100 * value:.1f}%"


def load_data():
    return {
        "annual": pd.read_csv(RESULTS / "annual_load_complete_observed_raw.csv"),
        "volume": pd.read_csv(RESULTS / "annual_runoff_volume_complete_observed_raw.csv"),
        "cumulative": pd.read_csv(RESULTS / "cumulative_load_2011_2025_raw.csv"),
        "ct_relative": pd.read_csv(RESULTS / "ct_relative_summary_raw.csv"),
        "spearman": pd.read_csv(RESULTS / "temporal_spearman_raw.csv"),
        "nrmse": pd.read_csv(RESULTS / "rmse_nrmse_bayes_ml_reconstruction_loyo_raw.csv"),
        "ml_coverage": pd.read_csv(RESULTS / "loyo_interval_coverage_by_year_target_raw.csv"),
        "bayes_coverage": pd.read_csv(RESULTS / "bayes_interval_coverage_by_year_target_raw.csv"),
        "importance": pd.read_csv(RESULTS / "feature_importance_descriptive_noncausal_raw.csv"),
        "negative": pd.read_csv(RESULTS / "bayes_negative_draw_sensitivity_summary_raw.csv"),
        "negative_ct": pd.read_csv(
            RESULTS / "bayes_negative_draw_sensitivity_ct_relative_summary_raw.csv"
        ),
    }


def coverage_summary(frame, target):
    subset = frame[(frame["Target"] == target) & frame["IntervalCoverage"].notna()].copy()
    minimum = subset.loc[subset["IntervalCoverage"].idxmin()]
    maximum = subset.loc[subset["IntervalCoverage"].idxmax()]
    weighted = (subset["IntervalCoverage"] * subset["n"]).sum() / subset["n"].sum()
    return minimum, maximum, weighted


def annual_load_rows(frame, analyte):
    decimals = 1 if analyte == "TSS" else 3
    subset = frame[frame["Analyte"] == analyte].copy()
    rows = []
    for year in sorted(subset["Year"].unique()):
        row = [str(int(year))]
        for treatment in TREATMENT_ORDER:
            record = subset[
                (subset["Year"] == year) & (subset["Treatment"] == treatment)
            ].iloc[0]
            observed = (
                f"{record['ObservedAnnualLoad_kg']:.{decimals}f}"
                if bool(record["ObservedAnnualLoadComplete"])
                else "NA"
            )
            row.extend(
                [
                    observed,
                    f"{record['BayesPosteriorMedian_kg']:.{decimals}f}",
                    f"{record['MLPointTotal_kg']:.{decimals}f}",
                ]
            )
        rows.append(row)
    return rows


def annual_volume_rows(frame):
    rows = []
    for year in sorted(frame["Year"].unique()):
        row = [str(int(year))]
        for treatment in TREATMENT_ORDER:
            record = frame[
                (frame["Year"] == year) & (frame["Treatment"] == treatment)
            ].iloc[0]
            observed = (
                f"{record['ObservedAnnualVolume_kL']:.1f}"
                if bool(record["ObservedAnnualVolumeComplete"])
                else "NA"
            )
            row.extend(
                [
                    observed,
                    f"{record['BayesPosteriorMean_kL']:.1f}",
                    f"{record['MLPointTotal_kL']:.1f}",
                ]
            )
        rows.append(row)
    return rows


def cumulative_rows(cumulative, ct_relative):
    rows = []
    for analyte in ["TSS", "TP", "TN"]:
        for method, scenario in (
            ("Bayes", "model_only"),
            ("ML", "full_record_model_only"),
        ):
            subset = cumulative[
                (cumulative["Method"] == method)
                & (cumulative["Scenario"] == scenario)
                & (cumulative["Analyte"] == analyte)
            ]
            centers = subset.set_index("Treatment")["primary_center"]
            relative = ct_relative[
                (ct_relative["Method"] == method)
                & (ct_relative["Scenario"] == scenario)
                & (ct_relative["Analyte"] == analyte)
            ].set_index("ComparisonTreatment")["primary_center"]
            decimals = 1 if analyte == "TSS" else 3
            rows.append(
                [
                    analyte,
                    method,
                    f"{centers['CT']:.{decimals}f}",
                    f"{centers['MT']:.{decimals}f}",
                    f"{relative['MT']:.1f}%",
                    f"{centers['ST']:.{decimals}f}",
                    f"{relative['ST']:.1f}%",
                ]
            )
    return rows


def spearman_rows(frame):
    frame = frame.copy()
    frame["Analyte"] = pd.Categorical(frame["Analyte"], ANALYTE_ORDER, ordered=True)
    rows = []
    for analyte, group in frame.sort_values(["Analyte", "Treatment"]).groupby(
        "Analyte", observed=False
    ):
        values = [str(analyte)]
        for treatment in TREATMENT_ORDER:
            record = group[group["Treatment"] == treatment].iloc[0]
            marker = "*" if bool(record["significant_unadjusted_p_lt_0_05"]) else ""
            values.append(f"{record['rho']:.2f}{marker}")
        rows.append(values)
    return rows


def nrmse_value(frame, display, series):
    record = frame[
        (frame["ComparisonSeries"] == series)
        & (
            ((frame["Analyte"] == display) if display != "Volume" else (frame["Target"] == "Volume_L"))
        )
    ].iloc[0]
    return 100 * record["NRMSE_mean_observed"]


def interval_span(frame, analyte):
    subset = frame[
        (frame["Analyte"] == analyte)
        & (frame["MLPointTotal_kg"] > 0)
        & frame["MLLower95_kg"].notna()
        & frame["MLUpper95_kg"].notna()
    ]
    return ((subset["MLUpper95_kg"] - subset["MLLower95_kg"]) / subset["MLPointTotal_kg"]).median()


def build_document():
    data = load_data()
    doc = Document()
    configure_styles(doc)
    add_title(doc)
    add_review_note(doc)

    annual = data["annual"]
    volume = data["volume"]
    cumulative = data["cumulative"]
    ct_relative = data["ct_relative"]

    ml_con_min, ml_con_max, ml_con_overall = coverage_summary(
        data["ml_coverage"], "Result_mg_L"
    )
    ml_vol_min, ml_vol_max, ml_vol_overall = coverage_summary(data["ml_coverage"], "Volume_L")
    bayes_con_min, bayes_con_max, bayes_con_overall = coverage_summary(
        data["bayes_coverage"], "Result_mg_L"
    )
    bayes_vol_min, bayes_vol_max, bayes_vol_overall = coverage_summary(
        data["bayes_coverage"], "Volume_L"
    )

    complete_main = annual[
        annual["Analyte"].isin(["TSS", "TP", "TN"]) & annual["ObservedAnnualLoadComplete"]
    ]
    n_complete_main = len(complete_main)
    n_possible_main = len(annual[annual["Analyte"].isin(["TSS", "TP", "TN"])])
    n_complete_volume = int(volume["ObservedAnnualVolumeComplete"].sum())

    doc.add_heading("3 Results and Discussion", level=1)
    add_body(
        doc,
        "Results are presented as evidence of ML calibration and predictor structure, annual and "
        "cumulative reconstruction, goodness of fit and temporal agreement, and cross-framework "
        "complementarity. Observed annual loads and runoff volumes are reported only when every "
        "expected physical event in a year-treatment group had the corresponding observation; "
        "incomplete annual subtotals are retained in the audit files but are not presented as "
        "observed annual values."
    )

    doc.add_heading("3.1 ML calibration, interval coverage, and predictor structure", level=2)
    add_body(
        doc,
        f"Year-specific interval coverage varied across the record (Figure 1). Bayesian "
        f"posterior-predictive concentration coverage ranged from {pct(bayes_con_min['IntervalCoverage'])} "
        f"in {int(bayes_con_min['Year'])} to {pct(bayes_con_max['IntervalCoverage'])} in "
        f"{int(bayes_con_max['Year'])}, with weighted overall coverage of "
        f"{pct(bayes_con_overall)}. ML outer-LOYO concentration coverage ranged from "
        f"{pct(ml_con_min['IntervalCoverage'])} in {int(ml_con_min['Year'])} to "
        f"{pct(ml_con_max['IntervalCoverage'])} in {int(ml_con_max['Year'])}, with weighted "
        f"overall coverage of {pct(ml_con_overall)}. Neither framework had concentration "
        f"observations available for annual coverage evaluation in 2018 or 2019."
    )
    add_body(
        doc,
        f"Runoff-volume coverage separated the frameworks more clearly. Bayesian "
        f"posterior-predictive volume coverage was {pct(bayes_vol_overall)} overall and ranged "
        f"from {pct(bayes_vol_min['IntervalCoverage'])} to {pct(bayes_vol_max['IntervalCoverage'])} "
        f"among evaluable years. ML outer-LOYO volume coverage was {pct(ml_vol_overall)} overall, "
        f"but fell to {pct(ml_vol_min['IntervalCoverage'])} in {int(ml_vol_min['Year'])}. "
        f"Coverage returned to {pct(ml_vol_max['IntervalCoverage'])} in the strongest year and "
        f"was high in 2023 and 2025, so the 2024 failure was not explained by a simple late-record "
        f"decline in data support. Instead, the 2024 result indicates a year-specific distribution "
        f"shift or event and measurement conditions not represented adequately by the remaining years."
    )
    add_body(
        doc,
        "This temporal instability is consistent with long-term edge-of-field monitoring studies "
        "that treat missing observations, changing measurement methods, and measurement uncertainty "
        "as structural features of operational datasets rather than minor data-cleaning problems "
        "(Harmel et al., 2006, 2023; Daniels et al., 2018; Harmel, King, et al., 2018).",
        highlight=True,
    )
    add_figure(
        doc,
        1,
        FIGURES / "postprocessing" / "loyo_interval_coverage_by_year.png",
        "Year-specific concentration and runoff-volume interval coverage for Bayesian "
        "posterior-predictive intervals and ML outer leave-one-year-out (LOYO) prediction intervals. "
        "Years without observed targets have no coverage estimate.",
        "Two-panel line chart of annual interval coverage for concentrations and runoff volume, "
        "comparing Bayesian posterior-predictive and ML leave-one-year-out intervals.",
        6.35,
    )

    add_body(
        doc,
        "The concentration model was dominated by analyte identity, which accounted for a mean "
        "importance of 58.9, followed by inflow concentration at 11.6 (Figure 2). Laboratory "
        "reporting fields, including method detection and reporting limits, formed the next tier of "
        "predictors. Thus, the concentration model primarily learned differences among analytes and "
        "source-water chemistry, while also using information associated with laboratory methods. "
        "The presence of laboratory metadata among influential variables supports prediction within "
        "the historical record but also shows that part of the learned structure was specific to the "
        "measurement system."
    )
    add_figure(
        doc,
        2,
        FIGURES / "postprocessing" / "feature_importance_concentration.png",
        "Descriptive CatBoost feature importance for concentration reconstruction. Bars show mean "
        "importance across fitted models and error bars show the corresponding standard deviation. "
        "Importance is descriptive and does not identify causal effects.",
        "Horizontal bar chart showing analyte identity and inflow concentration as the dominant "
        "features in the ML concentration model.",
        6.25,
    )

    add_body(
        doc,
        "Runoff-volume importance was distributed across seasonal timing and measurement context "
        "(Figure 3). Days until harvest had the largest mean importance (18.5), followed by days "
        "since planting (10.5), measurement method (9.9), day of year (9.9), flume method (8.2), "
        "and inflow volume (6.7). This broader distribution is consistent with runoff volume being "
        "conditioned by within-season timing, water delivery, and the measurement configuration. "
        "Treatment variables were not dominant, so feature importance did not provide direct "
        "evidence that tillage caused the reconstructed volume differences."
    )
    add_figure(
        doc,
        3,
        FIGURES / "postprocessing" / "feature_importance_volume.png",
        "Descriptive CatBoost feature importance for runoff-volume reconstruction. Bars show mean "
        "importance across fitted models and error bars show the corresponding standard deviation.",
        "Horizontal bar chart showing seasonal timing and measurement-method variables as the "
        "largest features in the ML runoff-volume model.",
        6.25,
    )

    doc.add_heading("3.2 Annual and cumulative reconstruction", level=2)
    add_body(
        doc,
        f"Only {n_complete_main} of {n_possible_main} possible annual load groups for TSS, TP, "
        f"and TN had a complete observed event record. TSS and TP each had eight complete "
        f"year-treatment groups, whereas TN had three, all in 2016. Complete annual runoff volume "
        f"was available for {n_complete_volume} of {len(volume)} year-treatment groups. Reporting "
        f"the remaining partial subtotals as annual observations would have implied that missing "
        f"events produced zero load or volume. Their exclusion makes the information gap explicit "
        f"and shows why model-based reconstruction was required for comparison across 2011-2025."
    )

    add_landscape_section(doc)
    add_body(
        doc,
        "Bayesian annual runoff-volume estimates were generally larger and less variable from year "
        "to year than the ML point totals (Figure 4; Table 1). Across the nine complete annual "
        "volume groups, the ML point totals were descriptively closer to the observed totals than "
        "the Bayesian posterior means, but this comparison covers only 2015-2017 and is not an "
        "independent validation of the reconstructed record. The study-period sums of annual "
        "central estimates were 5,787, 5,432, and 5,430 kL for Bayesian CT, MT, and ST, compared "
        "with 3,819, 3,481, and 3,512 kL for ML. Differences in reconstructed volume therefore "
        "propagated directly into differences in constituent loads."
    )
    add_figure(
        doc,
        4,
        FIGURES
        / "annual_complete_observed"
        / "annual_runoff_volume_complete_observed_v3p0.png",
        "Annual runoff volume by treatment from the Bayesian and ML reconstructions. Observed "
        "annual markers are included only when every expected physical event had a runoff-volume "
        "observation. ML annual prediction intervals are omitted from the primary panel for readability.",
        "Three-panel time-series chart of annual runoff volume for CT, MT, and ST, with Bayesian "
        "intervals, ML point totals, and complete observed annual markers.",
        9.45,
    )
    add_table_caption(doc, 1, "Annual runoff-volume reconstruction and complete observed totals (kL).")
    add_table(
        doc,
        [
            "Year",
            "CT\nObserved",
            "CT\nBayes",
            "CT\nML",
            "MT\nObserved",
            "MT\nBayes",
            "MT\nML",
            "ST\nObserved",
            "ST\nBayes",
            "ST\nML",
        ],
        annual_volume_rows(volume),
        [0.55] + [0.98] * 9,
        font_size=7,
    )
    add_table_note(
        doc,
        "Bayes is the posterior mean; ML is the physical-event point total. NA indicates that at "
        "least one expected physical event lacked an observed runoff-volume value.",
        RESULTS / "annual_runoff_volume_complete_observed_raw.csv",
    )

    add_landscape_section(doc)
    add_body(
        doc,
        "TSS produced the largest absolute loads and the strongest agreement in treatment direction "
        "(Figure 5; Table 2). The Bayesian reconstruction estimated a CT maximum of 1,136 kg in "
        "2018, whereas the ML reconstruction estimated a CT maximum of 1,014 kg in 2017. The "
        "complete 2017 CT observation was 2,163 kg, showing that both annual central estimates "
        "underrepresented that observed peak. Despite differences in annual magnitude, both "
        "frameworks produced lower study-period TSS totals for MT and ST than for CT."
    )
    add_figure(
        doc,
        5,
        FIGURES / "annual_complete_observed" / "annual_load_tss_complete_observed_v3p0.png",
        "Annual TSS load by treatment from the Bayesian and ML reconstructions. Observed annual "
        "markers are included only when all expected physical-event TSS loads were observed.",
        "Three-panel time-series chart of annual TSS loads for CT, MT, and ST, with Bayesian "
        "intervals, ML point totals, and complete observed annual markers.",
        9.45,
    )
    add_table_caption(doc, 2, "Annual TSS reconstruction and complete observed totals (kg).")
    add_table(
        doc,
        [
            "Year",
            "CT\nObserved",
            "CT\nBayes",
            "CT\nML",
            "MT\nObserved",
            "MT\nBayes",
            "MT\nML",
            "ST\nObserved",
            "ST\nBayes",
            "ST\nML",
        ],
        annual_load_rows(annual, "TSS"),
        [0.55] + [0.98] * 9,
        font_size=7,
    )
    add_table_note(
        doc,
        "Bayes is the posterior median; ML is the physical-event point total. NA denotes an "
        "incomplete observed annual event record. Raw negative Bayesian annual medians are retained.",
        RESULTS / "annual_load_complete_observed_raw.csv",
    )

    add_landscape_section(doc)
    add_body(
        doc,
        "TP showed the clearest disagreement in treatment interpretation (Figure 6; Table 3). Both "
        "frameworks identified 2025 as the largest reconstructed TP year in all three treatments. "
        "However, the Bayesian cumulative medians were 3.439, 2.279, and 2.092 kg for CT, MT, and "
        "ST, corresponding to 35.1% and 40.3% lower loads under MT and ST than CT. ML cumulative "
        "point totals were 3.053, 3.368, and 3.407 kg and therefore placed MT and ST 10.3% and "
        "11.6% above CT. TP conclusions were consequently sensitive to model structure rather than "
        "being common to both reconstructions."
    )
    add_figure(
        doc,
        6,
        FIGURES / "annual_complete_observed" / "annual_load_tp_complete_observed_v3p0.png",
        "Annual TP load by treatment from the Bayesian and ML reconstructions. Observed annual "
        "markers are included only when all expected physical-event TP loads were observed.",
        "Three-panel time-series chart of annual TP loads for CT, MT, and ST, with Bayesian "
        "intervals, ML point totals, and complete observed annual markers.",
        9.45,
    )
    add_table_caption(doc, 3, "Annual TP reconstruction and complete observed totals (kg).")
    add_table(
        doc,
        [
            "Year",
            "CT\nObserved",
            "CT\nBayes",
            "CT\nML",
            "MT\nObserved",
            "MT\nBayes",
            "MT\nML",
            "ST\nObserved",
            "ST\nBayes",
            "ST\nML",
        ],
        annual_load_rows(annual, "TP"),
        [0.55] + [0.98] * 9,
        font_size=7,
    )
    add_table_note(
        doc,
        "Bayes is the posterior median; ML is the physical-event point total. NA denotes an "
        "incomplete observed annual event record. Raw negative Bayesian annual medians are retained.",
        RESULTS / "annual_load_complete_observed_raw.csv",
    )

    add_landscape_section(doc)
    add_body(
        doc,
        "TN treatment direction was similar between frameworks, but reconstructed magnitude was "
        "not (Figure 7; Table 4). Bayesian cumulative medians were 32.463, 28.540, and 28.401 kg "
        "for CT, MT, and ST; ML point totals were 11.097, 7.850, and 7.965 kg. Thus, both methods "
        "placed MT and ST below CT, but Bayesian totals were approximately 2.9-3.6 times the ML "
        "totals. Agreement in treatment ordering therefore did not imply agreement in absolute load."
    )
    add_figure(
        doc,
        7,
        FIGURES / "annual_complete_observed" / "annual_load_tn_complete_observed_v3p0.png",
        "Annual TN load by treatment from the Bayesian and ML reconstructions. Observed annual "
        "markers are included only when all expected physical-event TN loads were observed.",
        "Three-panel time-series chart of annual TN loads for CT, MT, and ST, with Bayesian "
        "intervals, ML point totals, and complete observed annual markers.",
        9.45,
    )
    add_table_caption(doc, 4, "Annual TN reconstruction and complete observed totals (kg).")
    add_table(
        doc,
        [
            "Year",
            "CT\nObserved",
            "CT\nBayes",
            "CT\nML",
            "MT\nObserved",
            "MT\nBayes",
            "MT\nML",
            "ST\nObserved",
            "ST\nBayes",
            "ST\nML",
        ],
        annual_load_rows(annual, "TN"),
        [0.55] + [0.98] * 9,
        font_size=7,
    )
    add_table_note(
        doc,
        "Bayes is the posterior median; ML is the physical-event point total. NA denotes an "
        "incomplete observed annual event record.",
        RESULTS / "annual_load_complete_observed_raw.csv",
    )

    add_table_caption(
        doc,
        5,
        "Study-period cumulative load central estimates and treatment differences relative to CT.",
    )
    add_table(
        doc,
        [
            "Analyte",
            "Method",
            "CT\n(kg)",
            "MT\n(kg)",
            "MT lower\nthan CT",
            "ST\n(kg)",
            "ST lower\nthan CT",
        ],
        cumulative_rows(cumulative, ct_relative),
        [0.75, 0.85, 1.25, 1.25, 1.25, 1.25, 1.25],
        font_size=8,
    )
    add_table_note(
        doc,
        "Bayesian values are posterior medians from the model-only scenario. ML values are "
        "physical-event point totals from the full-record model-only scenario. Positive percentages "
        "indicate lower load than CT; negative percentages indicate higher load than CT.",
        RESULTS / "cumulative_load_2011_2025_raw.csv",
    )
    add_body(
        doc,
        "The shared TSS treatment direction, together with the smaller and less consistent TP and "
        "TN contrasts, follows the broader conservation-tillage literature in which sediment and "
        "particulate-associated nutrient responses are often more pronounced than dissolved "
        "constituent responses (Bjorneberg, Aase, et al., 2006; Kleinman et al., 2015; Prasad, "
        "Thomason, et al., 2023; Trimarco et al., 2025).",
        highlight=True,
    )

    doc.add_heading("3.3 Goodness of fit, temporal agreement, and uncertainty", level=2)
    nrmse = data["nrmse"]
    nrmse_text = {
        analyte: {
            "bayes": nrmse_value(nrmse, analyte, "Bayes posterior-predictive fit"),
            "ml": nrmse_value(nrmse, analyte, "ML full-record reconstruction"),
            "loyo": nrmse_value(nrmse, analyte, "ML outer-LOYO validation"),
        }
        for analyte in ["TSS", "TP", "TN", "Volume"]
    }
    add_body(
        doc,
        "The three error summaries answered different questions and were not directly interchangeable "
        "(Figure 8). Bayesian posterior-predictive error measured fit within the single Bayesian "
        "model fitted to the complete record; no Bayesian LOYO analysis was conducted. ML full-record "
        "error described the fitted point predictions used for historical reconstruction. ML outer-LOYO "
        "error measured prediction in a year excluded from model fitting and was therefore the strictest "
        "test of temporal transfer."
    )
    add_body(
        doc,
        f"Outer-LOYO NRMSE exceeded full-record ML NRMSE for every displayed target. For TSS, the "
        f"Bayesian fit, ML reconstruction, and ML LOYO NRMSE values were "
        f"{nrmse_text['TSS']['bayes']:.1f}%, {nrmse_text['TSS']['ml']:.1f}%, and "
        f"{nrmse_text['TSS']['loyo']:.1f}%, respectively. The corresponding values were "
        f"{nrmse_text['TP']['bayes']:.1f}%, {nrmse_text['TP']['ml']:.1f}%, and "
        f"{nrmse_text['TP']['loyo']:.1f}% for TP; {nrmse_text['TN']['bayes']:.1f}%, "
        f"{nrmse_text['TN']['ml']:.1f}%, and {nrmse_text['TN']['loyo']:.1f}% for TN; and "
        f"{nrmse_text['Volume']['bayes']:.1f}%, {nrmse_text['Volume']['ml']:.1f}%, and "
        f"{nrmse_text['Volume']['loyo']:.1f}% for runoff volume. Selenium produced the largest "
        f"normalized error because its small mean observed concentration magnified the NRMSE denominator."
    )
    add_body(
        doc,
        "RMSE and NRMSE are useful diagnostics but do not provide a model-independent definition of "
        "adequacy. Their interpretation depends on response scale, measurement error, benchmark "
        "definition, and intended model use; normalized error can become unstable for low-mean or "
        "episodic constituents (Moriasi et al., 2007; Harmel and Smith, 2007; Guzman et al., 2015; "
        "Yen et al., 2014; Beven, 2006).",
        highlight=True,
    )
    add_figure(
        doc,
        8,
        FIGURES / "gof_nrmse_bayes_ml_reconstruction_loyo_overall.png",
        "NRMSE for Bayesian posterior-predictive fit, ML full-record physical-event "
        "reconstruction, and ML outer-LOYO validation. Values are normalized by the mean observed "
        "target within each evaluation set. Selenium is shown with a capped display scale and an "
        "annotation for the full value.",
        "Grouped bar chart comparing Bayesian fit, ML reconstruction, and ML leave-one-year-out "
        "NRMSE for each analyte and runoff volume.",
        9.4,
    )

    add_table_caption(
        doc,
        6,
        "Spearman rank correlation between Bayesian and ML annual central estimates by treatment.",
    )
    add_table(
        doc,
        ["Analyte", "CT rho", "MT rho", "ST rho"],
        spearman_rows(data["spearman"]),
        [1.25, 1.25, 1.25, 1.25],
        font_size=9,
    )
    add_table_note(
        doc,
        "All correlations used 15 paired annual central estimates (2011-2025). * denotes an "
        "unadjusted p < 0.05. The tests are exploratory and were not adjusted for multiplicity. "
        "Correlation measures agreement in year ranking, not agreement in load magnitude.",
        RESULTS / "temporal_spearman_raw.csv",
    )
    add_body(
        doc,
        "Temporal rank agreement was strongest and most consistent for NO3, with rho values of "
        "0.55, 0.78, and 0.67 for CT, MT, and ST (Table 6). TN also had positive correlations in "
        "all treatments (0.51-0.64), even though the reconstructed TN magnitudes differed sharply. "
        "In contrast, TP correlations ranged from 0.11 to 0.43 and TSS correlations from 0.15 to "
        "0.37. These weak rank correlations indicate that agreement in cumulative TSS treatment "
        "direction arose despite different allocations of load among years."
    )
    tss_span = interval_span(annual, "TSS")
    tp_span = interval_span(annual, "TP")
    tn_span = interval_span(annual, "TN")
    add_body(
        doc,
        f"ML annual prediction intervals were too wide to display with the annual point series "
        f"without obscuring the treatment and year patterns. The median annual interval span was "
        f"{tss_span:.1f} times the TSS point total, {tp_span:.1f} times the TP point total, and "
        f"{tn_span:.1f} times the TN point total. The wide bands resulted from propagating "
        f"held-out residual uncertainty through event-level loads and annual sums; they are "
        f"prediction intervals, not confidence intervals for a fitted mean. Event-level conformal "
        f"coverage also does not guarantee nominal coverage for sums of correlated event loads. "
        f"The complete interval products remain in the supplementary outputs even though the "
        f"primary annual figures retain only ML point totals."
    )

    add_portrait_section(doc)
    doc.add_heading("3.4 Agreement, disagreement, and complementarity", level=2)
    add_body(
        doc,
        "The two reconstructions agreed most clearly on the direction of the TSS treatment contrast "
        "and on the ordering of TN treatments. They disagreed on the size and timing of annual TSS "
        "loads, the absolute magnitude of TN, and the direction of the cumulative TP contrast. "
        "These differences were not random diagnostic noise: they identified conclusions that were "
        "dependent on model structure and therefore require more cautious interpretation."
    )
    add_body(
        doc,
        "The Bayesian model represented treatment, year, analyte, and process relationships through "
        "a joint probabilistic structure and propagated posterior uncertainty into annual and "
        "cumulative loads. The ML workflow flexibly learned nonlinear associations from the observed "
        "record and provided an explicit held-out-year assessment of temporal prediction. Agreement "
        "between these structurally different estimates increased confidence in a qualitative "
        "pattern, whereas disagreement located the parts of the reconstruction most sensitive to "
        "assumptions. The analyses were not combined by model averaging, and neither estimate was "
        "treated as an independent replicate of the observed system."
    )
    add_body(
        doc,
        "This complementary interpretation is consistent with recent water-quality modeling work "
        "that separates flexible predictive performance from process or causal structure rather "
        "than treating one as a substitute for the other (Badrudeen et al., 2026; Longchamps, "
        "Lanza, Wei, et al., 2025).",
        highlight=True,
    )
    add_body(
        doc,
        "The ML reconstruction is consequently appropriate for interpolation within the historical "
        "study domain represented in model fitting, not for extrapolation to new management regimes "
        "or future years. A model would need to be recalibrated as new years, instruments, or "
        "management conditions are introduced, and its outer-LOYO performance would need to be "
        "reassessed. The Bayesian process structure provides a more defensible starting point for "
        "scenario analysis or forecasting, but forecasting was not validated in this study and "
        "should not be inferred from posterior fit alone."
    )

    doc.add_heading("4 Limitations", level=1)
    doc.add_heading("4.1 Incomplete observations and shared-data dependence", level=2)
    add_body(
        doc,
        "The most direct limitation was the scarcity of complete observed annual references. Only "
        "19 of 135 TSS, TP, and TN year-analyte-treatment groups and 9 of 45 runoff-volume groups "
        "could be compared with a complete observed annual total. Those complete groups were "
        "concentrated in 2015-2017 and therefore did not independently test early or late portions "
        "of the reconstruction. Both models were also calibrated from the same monitoring record. "
        "Agreement between them reflects robustness to model structure within shared data; it does "
        "not constitute independent confirmation of the unobserved loads."
    )

    doc.add_heading("4.2 ML transfer and uncertainty", level=2)
    add_body(
        doc,
        "The ML analysis used one algorithm with prespecified tuning and did not compare alternative "
        "learners or conduct a separate hyperparameter search within every held-out year. Its "
        "feature importance values were descriptive, noncausal, and partly associated with "
        "laboratory and measurement methods that changed through time. The extreme width of the "
        "propagated annual intervals limited their usefulness for treatment-scale inference even "
        "when event-level coverage was acceptable. Point totals remain useful historical "
        "reconstructions, but they do not support forecasting beyond the observed covariate and "
        "management domain."
    )

    doc.add_heading("4.3 Bayesian validation and nonnegative support", level=2)
    negative = data["negative"]
    neg_raw = negative[negative["Scenario"] == "raw_annual_draws_display_floor_only"].set_index(
        ["Analyte", "Treatment"]
    )
    negative_ct = data["negative_ct"]
    neg_ct = negative_ct[
        (negative_ct["Scenario"] == "annual_draw_truncation_at_zero")
        & negative_ct["Analyte"].isin(["TSS", "TP"])
    ].set_index(["Analyte", "ComparisonTreatment"])
    add_body(
        doc,
        "The Bayesian fit was evaluated with posterior-predictive diagnostics but not with a "
        "held-out-year design comparable to ML outer-LOYO. Its lower-tail distributions also "
        "permitted physically impossible negative annual loads. Negative annual draws represented "
        f"{neg_raw.loc[('TSS', 'CT'), 'percent_negative_annual_draws']:.1f}%-"
        f"{neg_raw.loc[('TSS', 'MT'), 'percent_negative_annual_draws']:.1f}% of TSS draws across "
        f"treatments and {neg_raw.loc[('TP', 'CT'), 'percent_negative_annual_draws']:.1f}%-"
        f"{neg_raw.loc[('TP', 'ST'), 'percent_negative_annual_draws']:.1f}% of TP draws. "
        "The annual figures use a nonnegative display scale, but the raw summaries and manuscript "
        "tables retain the modeled values rather than silently replacing them."
    )
    add_body(
        doc,
        f"Truncating annual Bayesian draws at zero reduced the median TSS treatment reduction from "
        f"48.2% to {neg_ct.loc[('TSS', 'MT'), 'median']:.1f}% for MT and from 40.9% to "
        f"{neg_ct.loc[('TSS', 'ST'), 'median']:.1f}% for ST. For TP, the corresponding reductions "
        f"changed from 35.1% to {neg_ct.loc[('TP', 'MT'), 'median']:.1f}% and from 40.3% to "
        f"{neg_ct.loc[('TP', 'ST'), 'median']:.1f}%. The sensitivity of TSS and TP treatment "
        f"contrasts to nonnegative support is a material limitation of the current Bayesian "
        f"specification and should be resolved before these cumulative percentages are treated as "
        f"final inferential estimates.",
        highlight=True,
    )

    doc.add_heading("4.4 Metrics and generalizability", level=2)
    add_body(
        doc,
        "NRMSE was unstable for constituents with small observed means, and Spearman correlation "
        "captured annual rank agreement without measuring calibration or absolute agreement. "
        "Unadjusted correlation p-values were exploratory. Finally, the reconstructions describe "
        "one irrigated long-term tillage experiment with its specific soils, crop rotations, "
        "irrigation practices, instruments, and laboratory history. Transfer to other sites or "
        "management systems requires new calibration and validation rather than direct application "
        "of the reported annual loads."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
